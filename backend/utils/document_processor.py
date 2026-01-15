# utils/document_processor.py
import os
import io
from typing import Tuple
import PyPDF2
from docx import Document
import re

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file bytes"""
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        # Clean up the text
        text = re.sub(r'\s+', ' ', text.strip())
        return text
        
    except Exception as e:
        raise ValueError(f"Error processing PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file bytes"""
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        # Clean up the text
        text = re.sub(r'\s+', ' ', text.strip())
        return text
        
    except Exception as e:
        raise ValueError(f"Error processing DOCX: {str(e)}")

def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from TXT file bytes"""
    try:
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                # Clean up the text
                text = re.sub(r'\s+', ' ', text.strip())
                return text
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode text file with any supported encoding")
        
    except Exception as e:
        raise ValueError(f"Error processing TXT: {str(e)}")

def process_uploaded_document(file_content: bytes, filename: str) -> Tuple[str, int]:
    """
    Process uploaded document and extract text
    Returns: (extracted_text, word_count)
    """
    try:
        # Get file extension
        file_ext = os.path.splitext(filename.lower())[1]
        
        if file_ext == '.pdf':
            text = extract_text_from_pdf(file_content)
        elif file_ext == '.docx':
            text = extract_text_from_docx(file_content)
        elif file_ext in ['.txt', '.text']:
            text = extract_text_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}. Supported types: PDF, DOCX, TXT")
        
        if not text or len(text.strip()) < 10:
            raise ValueError("No readable text found in the document")
        
        word_count = len(text.split())
        return text, word_count
        
    except Exception as e:
        raise ValueError(f"Document processing failed: {str(e)}")

def count_words(text: str) -> int:
    """Count words in text"""
    return len(text.split())

def validate_text_length(text: str, max_words: int = 5000) -> Tuple[bool, int, str]:
    """
    Validate text length
    Returns: (is_valid, word_count, status_message)
    """
    word_count = count_words(text)
    
    if word_count > max_words:
        return False, word_count, f"Text exceeds {max_words} word limit"
    elif word_count < 50:
        return False, word_count, "Text must be at least 50 words"
    else:
        return True, word_count, "Text length is valid"